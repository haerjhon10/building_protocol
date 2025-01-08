import streamlit as st
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import base64

DEFAULT_TESTS = [
    "ריצוף",
    "חזות הגמר והצביעה",
    "אלומיניום",
    "חשמל ותאורה"
]

def create_inspection_protocol(title, sections):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title_element = doc.add_heading(title, 0)
    title_element.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_section(heading, rows):
        doc.add_heading(heading, level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table = doc.add_table(rows=len(rows)+1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        widths = [Cm(4), Cm(4), Cm(4), Cm(8)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        headers = ['אישור לקוח לתיקון הערות',
                   'הערות',
                   'תקין ומתפקד',
                   'תיאור הבדיקה']
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(header)
            run.bold = True
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w'))))

        for i, row_text in enumerate(rows, 1):
            cell = table.cell(i, 3)
            cell.text = row_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for j in [0, 1, 2]:
                cell = table.cell(i, j)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    for heading, rows in sections.items():
        add_section(heading, rows)

    file_path = 'inspection_protocol.docx'
    doc.save(file_path)
    return file_path

st.title("Dynamic Inspection Protocol Generator")

if 'sections' not in st.session_state:
    st.session_state.sections = {}

title = st.text_input("Document Title", "פרוטוקול מסירת מבנה בית כנסת - בדיקות נדרשות")

# Space management
col1, col2, col3 = st.columns([0.6, 0.2, 0.2])
with col1:
    new_section_name = st.text_input("Add a New Space (Section Name)")
with col2:
    st.write("")
    if st.button("Add Space", use_container_width=True) and new_section_name:
        st.session_state.sections[new_section_name] = DEFAULT_TESTS.copy()
        st.rerun()

# Display and manage sections
for section_name in list(st.session_state.sections.keys()):
    st.subheader(section_name)
    
    if st.button(f"Delete Space {section_name}"):
        del st.session_state.sections[section_name]
        st.rerun()
        
    # Test management
    test_input = st.text_input(f"Add a test for {section_name}", key=f"test_{section_name}")
    if st.button(f"Add Test to {section_name}", key=f"add_test_{section_name}"):
        if test_input:
            st.session_state.sections[section_name].append(test_input)
            st.rerun()

    # Display tests
    for idx, test in enumerate(st.session_state.sections[section_name]):
        cols = st.columns([0.9, 0.1])
        with cols[0]:
            st.write(f"- {test}")
        with cols[1]:
            if st.button("Delete", key=f"delete_{section_name}_{idx}"):
                st.session_state.sections[section_name].pop(idx)
                st.rerun()

if st.button("Generate Document"):
    file_path = create_inspection_protocol(title, st.session_state.sections)
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="inspection_protocol.docx">Download Document</a>'
        st.markdown(href, unsafe_allow_html=True)
